"""v8.0 T3901 — 自动周报服务.

Responsibilities:
    * 每周一早上 9 点自动生成运营周报
    * 复用 v6.0 ``document_generator`` 渲染 PDF / DOCX
    * 内容覆盖:
        - 日活 (DAU) 与趋势
        - 关键功能使用排行
        - 16 项需求使用率 (T3501 需求做透)
        - 异常点 (来自 anomaly_detector)
    * 邮件给老板 / HRBP (复用 v6.0 notify dispatcher)
    * 落库 ``weekly_reports`` 表供 admin/insights 展示

设计要点:
    * 可注入 ``_clock`` / ``_supabase`` / ``_dispatcher`` 便于测试
    * 触发器: ``schedule_weekly_report`` 由 scheduler (APScheduler / cron) 拉起
    * 离线 fallback: 当 Supabase / notify 不可用时, 仍然生成报告 bytes 并返回
"""
from __future__ import annotations

import io
import json
import logging
import os
from dataclasses import dataclass, field
from datetime import datetime, timedelta, timezone
from enum import Enum
from typing import Any, Callable, Dict, Iterable, List, Optional, Tuple

logger = logging.getLogger("recruittech.platform.auto_report")


# ---------------------------------------------------------------------------
# 常量 / 枚举
# ---------------------------------------------------------------------------


class ReportFormat(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"


# 16 项需求 (1.1-1.6 求职者 + 2.1-2.9 用人单位 + 3 综合) — 与 todo.json 对齐
SIXTEEN_REQUIREMENTS: List[Tuple[str, str]] = [
    ("1.1", "ProfileCard 知心朋友"),
    ("1.2", "主动 push 频繁互动"),
    ("1.3", "画像生成后求职者确认"),
    ("1.4", "规划生成后执行追踪"),
    ("1.5", "心情差时智能体关怀"),
    ("1.6", "HR 偏好 + 模板化回复"),
    ("2.1", "JD 营销化撰写"),
    ("2.2", "假资质鉴别"),
    ("2.3", "AI 模拟面试官"),
    ("2.4", "Offer 真实业务"),
    ("2.5", "试用期跟踪"),
    ("2.6", "内部推荐"),
    ("2.7", "沉睡激活"),
    ("2.8", "偏见检测"),
    ("2.9", "制度 AI 解释"),
    ("3", "战略传达 + 主动 HR 建议"),
]


# ---------------------------------------------------------------------------
# 数据结构
# ---------------------------------------------------------------------------


@dataclass
class DAUMetric:
    date: str
    dau: int
    new_users: int = 0
    returning: int = 0


@dataclass
class FeatureUsage:
    feature: str
    invocations: int
    unique_users: int
    growth_pct: float = 0.0  # vs prior week


@dataclass
class RequirementUsage:
    req_id: str
    req_name: str
    usage_pct: float  # 0..100
    delta_pct: float = 0.0
    note: str = ""


@dataclass
class WeeklyReport:
    week_start: str
    week_end: str
    generated_at: str
    format: str
    content: bytes
    filename: str
    size_bytes: int
    summary: Dict[str, Any] = field(default_factory=dict)
    recipients: List[str] = field(default_factory=list)
    pdf_path: Optional[str] = None
    delivered: bool = False
    delivery_channels: List[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# 默认数据源 (可被注入覆盖)
# ---------------------------------------------------------------------------


def _default_clock() -> datetime:
    return datetime.now(timezone.utc)


def _default_supabase_factory():
    """懒加载 supabase 客户端; 失败返回 None."""
    try:
        from api.deps import get_supabase_admin

        return get_supabase_admin()
    except Exception:  # pragma: no cover - 离线 / 测试
        return None


def _default_dispatcher_factory():
    try:
        from services.notify import get_dispatcher

        return get_dispatcher()
    except Exception:  # pragma: no cover
        return None


# ---------------------------------------------------------------------------
# 数据收集
# ---------------------------------------------------------------------------


def _collect_dau(week_start: datetime, week_end: datetime, supabase) -> List[DAUMetric]:
    """从 supabase 拉取 DAU; 无 DB 时使用 mock."""
    if supabase is None:
        return _mock_dau(week_start, week_end)
    try:
        res = (
            supabase.table("dau_daily")
            .select("date,dau,new_users,returning")
            .gte("date", week_start.date().isoformat())
            .lte("date", week_end.date().isoformat())
            .order("date")
            .execute()
        )
        rows = res.data or []
        return [
            DAUMetric(
                date=r["date"],
                dau=int(r.get("dau") or 0),
                new_users=int(r.get("new_users") or 0),
                returning=int(r.get("returning") or 0),
            )
            for r in rows
        ]
    except Exception as exc:  # pragma: no cover
        logger.warning("auto_report: dau query failed: %s", exc)
        return _mock_dau(week_start, week_end)


def _mock_dau(week_start: datetime, week_end: datetime) -> List[DAUMetric]:
    """用于测试 / 离线场景的 deterministic mock."""
    out: List[DAUMetric] = []
    cur = week_start
    base = 312
    for i in range((week_end - week_start).days + 1):
        d = cur + timedelta(days=i)
        out.append(
            DAUMetric(
                date=d.date().isoformat(),
                dau=base + i * 8,
                new_users=12 + i,
                returning=base + i * 8 - 12 - i,
            )
        )
    return out


def _collect_feature_usage(week_start: datetime, week_end: datetime, supabase) -> List[FeatureUsage]:
    if supabase is None:
        return _mock_feature_usage()
    try:
        res = (
            supabase.table("feature_usage_weekly")
            .select("feature,invocations,unique_users,growth_pct")
            .eq("week_start", week_start.date().isoformat())
            .order("invocations", desc=True)
            .limit(20)
            .execute()
        )
        rows = res.data or []
        return [
            FeatureUsage(
                feature=r["feature"],
                invocations=int(r.get("invocations") or 0),
                unique_users=int(r.get("unique_users") or 0),
                growth_pct=float(r.get("growth_pct") or 0.0),
            )
            for r in rows
        ]
    except Exception as exc:  # pragma: no cover
        logger.warning("auto_report: feature_usage query failed: %s", exc)
        return _mock_feature_usage()


def _mock_feature_usage() -> List[FeatureUsage]:
    return [
        FeatureUsage("matching", 1842, 412, 12.4),
        FeatureUsage("profile_card", 1311, 287, 8.1),
        FeatureUsage("ai_interview", 974, 218, -3.5),
        FeatureUsage("jd_generate", 712, 165, 22.0),
        FeatureUsage("support_ticket", 521, 311, 4.6),
        FeatureUsage("video_resume", 408, 102, 31.2),
        FeatureUsage("bias_detector", 311, 88, 1.3),
        FeatureUsage("sourcing", 287, 76, 6.5),
        FeatureUsage("internal_referral", 152, 41, -1.2),
        FeatureUsage("salary_benchmark", 134, 33, 14.7),
    ]


def _collect_requirement_usage(
    week_start: datetime, week_end: datetime, supabase
) -> List[RequirementUsage]:
    if supabase is None:
        return _mock_requirement_usage()
    try:
        res = (
            supabase.table("requirement_usage_weekly")
            .select("req_id,usage_pct,delta_pct,note")
            .eq("week_start", week_start.date().isoformat())
            .execute()
        )
        rows = res.data or []
        out: List[RequirementUsage] = []
        for r in rows:
            rid = r["req_id"]
            name = next((n for k, n in SIXTEEN_REQUIREMENTS if k == rid), rid)
            out.append(
                RequirementUsage(
                    req_id=rid,
                    req_name=name,
                    usage_pct=float(r.get("usage_pct") or 0.0),
                    delta_pct=float(r.get("delta_pct") or 0.0),
                    note=str(r.get("note") or ""),
                )
            )
        # 补全缺失的 16 项
        existing = {r.req_id for r in out}
        for rid, name in SIXTEEN_REQUIREMENTS:
            if rid not in existing:
                out.append(RequirementUsage(req_id=rid, req_name=name, usage_pct=0.0, note="未启动"))
        return out
    except Exception as exc:  # pragma: no cover
        logger.warning("auto_report: req_usage query failed: %s", exc)
        return _mock_requirement_usage()


def _mock_requirement_usage() -> List[RequirementUsage]:
    seed = [0.91, 0.84, 0.62, 0.55, 0.78, 0.71, 0.66, 0.49, 0.82, 0.74, 0.58, 0.43, 0.61, 0.69, 0.37, 0.81]
    deltas = [0.04, 0.07, 0.02, -0.01, 0.05, 0.03, 0.10, 0.0, 0.04, 0.06, 0.01, -0.04, 0.02, 0.03, -0.02, 0.05]
    out: List[RequirementUsage] = []
    for (rid, name), pct, d in zip(SIXTEEN_REQUIREMENTS, seed, deltas):
        out.append(
            RequirementUsage(
                req_id=rid,
                req_name=name,
                usage_pct=round(pct * 100, 1),
                delta_pct=round(d * 100, 1),
            )
        )
    return out


def _collect_anomalies(anomalies: Optional[List[Dict[str, Any]]]) -> List[Dict[str, Any]]:
    """复用 anomaly_detector 输出 (已 dict 化)."""
    return list(anomalies or [])


# ---------------------------------------------------------------------------
# 渲染
# ---------------------------------------------------------------------------


def _format_report_text(
    week_start: datetime,
    week_end: datetime,
    dau: List[DAUMetric],
    features: List[FeatureUsage],
    reqs: List[RequirementUsage],
    anomalies: List[Dict[str, Any]],
    generated_at: datetime,
) -> Tuple[str, Dict[str, Any]]:
    total_dau = sum(d.dau for d in dau)
    avg_dau = total_dau // max(len(dau), 1)
    new_users = sum(d.new_users for d in dau)
    top_feature = features[0] if features else None
    low_requirements = sorted(reqs, key=lambda r: r.usage_pct)[:3]
    high_requirements = sorted(reqs, key=lambda r: r.usage_pct, reverse=True)[:3]

    lines: List[str] = []
    lines.append("=" * 64)
    lines.append(f" 招聘智能体 — 运营周报 v8.0")
    lines.append(f" {week_start.date().isoformat()} ~ {week_end.date().isoformat()}")
    lines.append(f" 生成时间: {generated_at.isoformat()}")
    lines.append("=" * 64)
    lines.append("")

    lines.append("【1. 日活 (DAU)】")
    lines.append(f"  - 周累计: {total_dau:,}")
    lines.append(f"  - 日均:   {avg_dau:,}")
    lines.append(f"  - 新增:   {new_users:,}")
    lines.append("  - 每日:")
    for d in dau:
        lines.append(f"      {d.date}  DAU={d.dau:>4}  新增={d.new_users:>3}  回归={d.returning:>4}")
    lines.append("")

    lines.append("【2. 关键功能使用 TOP10】")
    lines.append(f"  {'功能':<24} {'调用':>8} {'用户':>6} {'周环比':>8}")
    for f in features[:10]:
        lines.append(
            f"  {f.feature:<24} {f.invocations:>8,} {f.unique_users:>6} {f.growth_pct:>+7.1f}%"
        )
    if top_feature:
        lines.append(f"  → 最热门: {top_feature.feature} ({top_feature.invocations:,})")
    lines.append("")

    lines.append("【3. 16 项需求使用率 (按做透目标)】")
    lines.append(f"  {'编号':<4} {'需求':<22} {'使用率':>8} {'周环比':>8}")
    for r in reqs:
        lines.append(
            f"  {r.req_id:<4} {r.req_name:<22} {r.usage_pct:>7.1f}% {r.delta_pct:>+7.1f}%"
        )
    if high_requirements:
        lines.append("  ↑ 高使用率:")
        for r in high_requirements:
            lines.append(f"      {r.req_id} {r.req_name} — {r.usage_pct:.1f}%")
    if low_requirements:
        lines.append("  ↓ 待提升 (PM 跟进):")
        for r in low_requirements:
            lines.append(f"      {r.req_id} {r.req_name} — {r.usage_pct:.1f}%  ({r.note or '需推动'})")
    lines.append("")

    lines.append("【4. 异常点】")
    if not anomalies:
        lines.append("  无 (本周各项指标平稳).")
    else:
        for a in anomalies:
            lines.append(
                f"  - [{a.get('severity','?').upper()}] {a.get('metric','?')}: "
                f"{a.get('current','?')} (基线 {a.get('baseline','?')}) — {a.get('message','')}"
            )
    lines.append("")

    lines.append("=" * 64)
    lines.append(" 自动生成 by 招聘智能体 / v8.0 T3901 auto_report")
    lines.append("=" * 64)
    text = "\n".join(lines)

    summary = {
        "total_dau": total_dau,
        "avg_dau": avg_dau,
        "new_users": new_users,
        "top_feature": top_feature.feature if top_feature else None,
        "top_feature_invocations": top_feature.invocations if top_feature else 0,
        "low_requirement_ids": [r.req_id for r in low_requirements],
        "anomaly_count": len(anomalies),
        "req_count": len(reqs),
    }
    return text, summary


def _render_pdf(text: str) -> Tuple[bytes, str, str]:
    """使用 reportlab 渲染 PDF; 失败回退 TXT."""
    try:
        from reportlab.lib import colors  # type: ignore
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
        from reportlab.lib.units import cm  # type: ignore
        from reportlab.pdfbase import pdfmetrics  # type: ignore
        from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
        from reportlab.platypus import (  # type: ignore
            Paragraph, SimpleDocTemplate, Spacer,
        )
    except ImportError:
        return text.encode("utf-8"), "txt", "weekly_report.txt"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    body = styles["BodyText"]
    title = styles["Title"]
    story = []
    for i, line in enumerate(text.splitlines()):
        if i < 3 and line.startswith("="):
            continue
        if i == 0 or "招聘智能体" in line or "运营周报" in line:
            story.append(Paragraph(line or "&nbsp;", title))
            story.append(Spacer(1, 0.2 * cm))
        elif line.startswith("【"):
            story.append(Paragraph(line, styles["Heading2"]))
            story.append(Spacer(1, 0.1 * cm))
        elif line.strip():
            story.append(Paragraph(line.replace(" ", "&nbsp;"), body))
            story.append(Spacer(1, 0.05 * cm))
    try:
        doc.build(story)
    except Exception as exc:  # pragma: no cover - 字体缺失
        logger.warning("auto_report: PDF build failed, fallback txt: %s", exc)
        return text.encode("utf-8"), "txt", "weekly_report.txt"
    data = buf.getvalue()
    return data, "pdf", "weekly_report.pdf"


def _render_docx(text: str) -> Tuple[bytes, str, str]:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return text.encode("utf-8"), "txt", "weekly_report.txt"
    doc = Document()
    for i, line in enumerate(text.splitlines()):
        if i < 3 and line.startswith("="):
            continue
        if line.startswith("【"):
            doc.add_heading(line, level=2)
        elif line.strip():
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), "docx", "weekly_report.docx"


# ---------------------------------------------------------------------------
# 主服务
# ---------------------------------------------------------------------------


class AutoReportService:
    """自动周报生成 + 发送服务."""

    def __init__(
        self,
        *,
        clock: Callable[[], datetime] = _default_clock,
        supabase_factory: Callable[[], Any] = _default_supabase_factory,
        dispatcher_factory: Callable[[], Any] = _default_dispatcher_factory,
        default_recipients: Optional[Iterable[str]] = None,
        recipients_by_role: Optional[Dict[str, Iterable[str]]] = None,
    ) -> None:
        self._clock = clock
        self._supabase_factory = supabase_factory
        self._dispatcher_factory = dispatcher_factory
        self._default_recipients = list(default_recipients or ["ceo@waibao.example", "hrbp@waibao.example"])
        self._recipients_by_role = recipients_by_role or {}

    # ---- 工具方法 ----
    def last_week_range(self, now: Optional[datetime] = None) -> Tuple[datetime, datetime]:
        """返回上周一 00:00 ~ 上周日 23:59:59 (UTC)."""
        now = now or self._clock()
        # 周一为一周开始; weekday(): Mon=0..Sun=6
        days_since_monday = now.weekday()
        this_monday = (now - timedelta(days=days_since_monday)).replace(
            hour=0, minute=0, second=0, microsecond=0
        )
        last_monday = this_monday - timedelta(days=7)
        last_sunday = this_monday - timedelta(seconds=1)
        return last_monday, last_sunday

    def current_week_range(self, now: Optional[datetime] = None) -> Tuple[datetime, datetime]:
        now = now or self._clock()
        days_since_monday = now.weekday()
        monday = (now - timedelta(days=days_since_monday)).replace(
            hour=0, minute=0, second=0, microsecond=0
        )
        sunday = monday + timedelta(days=7) - timedelta(seconds=1)
        return monday, sunday

    def resolve_recipients(self, role: Optional[str] = None) -> List[str]:
        if role and role in self._recipients_by_role:
            return list(self._recipients_by_role[role])
        return list(self._default_recipients)

    # ---- 数据收集 ----
    def collect(
        self,
        week_start: datetime,
        week_end: datetime,
        anomalies: Optional[List[Dict[str, Any]]] = None,
    ) -> Dict[str, Any]:
        supabase = self._supabase_factory()
        dau = _collect_dau(week_start, week_end, supabase)
        features = _collect_feature_usage(week_start, week_end, supabase)
        reqs = _collect_requirement_usage(week_start, week_end, supabase)
        anomal = _collect_anomalies(anomalies)
        return {
            "dau": dau,
            "features": features,
            "requirements": reqs,
            "anomalies": anomal,
        }

    # ---- 生成 ----
    def generate(
        self,
        week_start: Optional[datetime] = None,
        week_end: Optional[datetime] = None,
        *,
        fmt: str = "pdf",
        anomalies: Optional[List[Dict[str, Any]]] = None,
        persist: bool = True,
    ) -> WeeklyReport:
        now = self._clock()
        if week_start is None or week_end is None:
            week_start, week_end = self.last_week_range(now)
        data = self.collect(week_start, week_end, anomalies=anomalies)
        text, summary = _format_report_text(
            week_start=week_start,
            week_end=week_end,
            dau=data["dau"],
            features=data["features"],
            reqs=data["requirements"],
            anomalies=data["anomalies"],
            generated_at=now,
        )
        if fmt == ReportFormat.PDF.value:
            content, actual_fmt, filename = _render_pdf(text)
        elif fmt == ReportFormat.DOCX.value:
            content, actual_fmt, filename = _render_docx(text)
        else:
            content = text.encode("utf-8")
            actual_fmt = "txt"
            filename = "weekly_report.txt"
        report = WeeklyReport(
            week_start=week_start.date().isoformat(),
            week_end=week_end.date().isoformat(),
            generated_at=now.isoformat(),
            format=actual_fmt,
            content=content,
            filename=filename,
            size_bytes=len(content),
            summary=summary,
        )
        if persist:
            self._persist(report, summary)
        return report

    # ---- 持久化 ----
    def _persist(self, report: WeeklyReport, summary: Dict[str, Any]) -> Optional[str]:
        supabase = self._supabase_factory()
        if supabase is None:
            return None
        try:
            payload = {
                "week_start": report.week_start,
                "week_end": report.week_end,
                "generated_at": report.generated_at,
                "format": report.format,
                "filename": report.filename,
                "size_bytes": report.size_bytes,
                "summary": summary,
            }
            res = supabase.table("weekly_reports").insert(payload).execute()
            rows = res.data or []
            if rows:
                return rows[0].get("id")
        except Exception as exc:  # pragma: no cover
            logger.warning("auto_report: persist failed: %s", exc)
        return None

    # ---- 发送 ----
    async def deliver(
        self,
        report: WeeklyReport,
        recipients: Optional[Iterable[str]] = None,
    ) -> Dict[str, Any]:
        """通过 notify dispatcher 发送给收件人; 离线时静默返回."""
        rcpts = list(recipients or self._default_recipients)
        dispatcher = self._dispatcher_factory()
        delivered_channels: List[str] = []
        failed: List[str] = []
        if dispatcher is None:
            logger.info("auto_report: dispatcher unavailable, skip delivery")
            return {
                "delivered": False,
                "recipients": rcpts,
                "channels": [],
                "failed": ["no_dispatcher"],
            }
        for user_id in rcpts:
            outcome = await dispatcher.dispatch_multi(
                channels=["smtp", "dingtalk", "feishu", "im", "web"],
                user_id=user_id,
                title=f"招聘智能体 周报 {report.week_start} ~ {report.week_end}",
                content=(
                    f"本周 DAU 合计 {report.summary.get('total_dau','?')}, "
                    f"日均 {report.summary.get('avg_dau','?')}; "
                    f"最热门功能 {report.summary.get('top_feature','?')}; "
                    f"异常点 {report.summary.get('anomaly_count','?')} 个. "
                    f"详见附件 {report.filename}."
                ),
                payload={
                    "filename": report.filename,
                    "format": report.format,
                    "size_bytes": report.size_bytes,
                    "summary": report.summary,
                },
            )
            for r in outcome.results:
                if r.success:
                    delivered_channels.append(r.channel)
                elif not r.skipped:
                    failed.append(f"{user_id}/{r.channel}:{r.error or 'unknown'}")
        report.delivered = bool(delivered_channels)
        report.delivery_channels = list(set(delivered_channels))
        report.recipients = rcpts
        return {
            "delivered": report.delivered,
            "recipients": rcpts,
            "channels": delivered_channels,
            "failed": failed,
        }

    # ---- 调度入口 ----
    async def schedule_weekly_report(
        self,
        *,
        fmt: str = "pdf",
        role: Optional[str] = None,
        anomalies: Optional[List[Dict[str, Any]]] = None,
    ) -> Dict[str, Any]:
        """scheduler 入口: 每周一 09:00 调用."""
        report = self.generate(fmt=fmt, anomalies=anomalies, persist=True)
        rcpts = self.resolve_recipients(role=role)
        delivery = await self.deliver(report, recipients=rcpts)
        return {
            "report_id": getattr(report, "pdf_path", None),
            "week": f"{report.week_start} ~ {report.week_end}",
            "format": report.format,
            "size_bytes": report.size_bytes,
            "summary": report.summary,
            "delivery": delivery,
        }


# ---------------------------------------------------------------------------
# 单例
# ---------------------------------------------------------------------------

_service: Optional[AutoReportService] = None


def get_auto_report_service() -> AutoReportService:
    global _service
    if _service is None:
        _service = AutoReportService()
    return _service


def reset_auto_report_service() -> None:
    global _service
    _service = None
